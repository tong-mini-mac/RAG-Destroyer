from fpdf import FPDF
import os
import shutil
from docx import Document

class FileExporter:
    def __init__(self, output_dir="output"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        # Font path for Mac (Thai support)
        self.font_path = "/System/Library/Fonts/Supplemental/Ayuthaya.ttf"

    def export_to_docx(self, markdown_path, target_name=None):
        """Converts Markdown to a simple Docx file."""
        if not os.path.exists(markdown_path):
            return None
            
        with open(markdown_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # Simple conversion (strip YAML frontmatter)
        if content.startswith("---"):
            parts = content.split("---", 2)
            if len(parts) >= 3:
                content = parts[2].strip()

        doc = Document()
        for line in content.split("\n"):
            if line.startswith("# "):
                doc.add_heading(line[2:], 0)
            elif line.startswith("## "):
                doc.add_heading(line[3:], 1)
            elif line.startswith("### "):
                doc.add_heading(line[4:], 2)
            else:
                doc.add_paragraph(line)

        target_name = target_name or os.path.basename(markdown_path).replace(".md", ".docx")
        if not target_name.endswith(".docx"):
            target_name += ".docx"
            
        save_path = os.path.join(self.output_dir, target_name)
        doc.save(save_path)
        return save_path

    def export_to_pdf(self, markdown_path, target_name=None):
        """Converts Markdown to a professional PDF with Thai/English support."""
        if not os.path.exists(markdown_path):
            return None

        with open(markdown_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # Strip YAML
        if content.startswith("---"):
            parts = content.split("---", 2)
            if len(parts) >= 3:
                content = parts[2].strip()

        pdf = FPDF()
        pdf.add_page()
        
        # Add Thai Font
        if os.path.exists(self.font_path):
            pdf.add_font("Ayuthaya", "", self.font_path)
            pdf.set_font("Ayuthaya", size=11)
        else:
            pdf.set_font("Helvetica", size=11)

        # Header
        pdf.set_font(size=16)
        pdf.set_text_color(20, 40, 100)
        pdf.cell(0, 10, "RAG-Destroyer: Global Guru Insights", ln=True, align='C')
        pdf.ln(5)
        
        # Content
        pdf.set_font(size=11)
        pdf.set_text_color(0, 0, 0)
        
        for line in content.split("\n"):
            if line.startswith("# "):
                pdf.set_font(size=14)
                pdf.cell(0, 10, line[2:], ln=True)
                pdf.set_font(size=11)
            elif line.startswith("## "):
                pdf.set_font(size=12)
                pdf.cell(0, 8, line[3:], ln=True)
                pdf.set_font(size=11)
            else:
                pdf.multi_cell(0, 7, line)
        
        target_name = target_name or os.path.basename(markdown_path).replace(".md", ".pdf")
        if not target_name.endswith(".pdf"):
            target_name += ".pdf"
            
        save_path = os.path.join(self.output_dir, target_name)
        pdf.output(save_path)
        return save_path
